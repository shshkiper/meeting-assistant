"""Protocol read, edit and DOCX export."""

import uuid
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.api.deps import get_current_user
from app.core.database import get_db
from app.models.models import Meeting, Protocol, Transcript, User
from app.schemas.schemas import ProtocolRead, ProtocolUpdate

router = APIRouter()


@router.get("/{meeting_id}", response_model=ProtocolRead)
async def get_protocol(
    meeting_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Protocol).where(Protocol.meeting_id == meeting_id))
    protocol = result.scalar_one_or_none()
    if not protocol:
        raise HTTPException(status_code=404, detail="Protocol not ready yet")
    return protocol


@router.patch("/{meeting_id}", response_model=ProtocolRead)
async def update_protocol(
    meeting_id: uuid.UUID,
    body: ProtocolUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Protocol).where(Protocol.meeting_id == meeting_id))
    protocol = result.scalar_one_or_none()
    if not protocol:
        raise HTTPException(status_code=404, detail="Protocol not found")
    for field, value in body.model_dump(exclude_none=True).items():
        setattr(protocol, field, value)
    await db.commit()
    await db.refresh(protocol)
    return protocol


@router.get("/{meeting_id}/export/docx")
async def export_docx(
    meeting_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export protocol as a Word document."""
    result = await db.execute(select(Protocol).where(Protocol.meeting_id == meeting_id))
    protocol = result.scalar_one_or_none()
    if not protocol or not protocol.content_md:
        raise HTTPException(status_code=404, detail="Protocol not found")

    mtg_res = await db.execute(select(Meeting).where(Meeting.id == meeting_id))
    meeting = mtg_res.scalar_one()

    docx_bytes = _build_docx(protocol.content_md, meeting.title)
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="protocol_{meeting_id}.docx"'},
    )


def _build_docx(markdown_text: str, title: str) -> bytes:
    """Convert markdown protocol to DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import re

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in markdown_text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line), style="List Number")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        else:
            # Handle inline bold **text**
            p = doc.add_paragraph()
            parts = re.split(r"\*\*(.+?)\*\*", line)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
